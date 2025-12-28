"""
Document Parsing API Route - 真实解析版
"""
from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
import os
from pathlib import Path
import sys

router = APIRouter()

# 添加 rag single 路径
RAG_DIR = Path(__file__).parent.parent.parent.parent / "rag single"
UPLOAD_DIR = RAG_DIR / "uploads"

class ParseRequest(BaseModel):
    fileId: str

@router.post("/parse")
async def parse_document(request: ParseRequest):
    """解析文档，返回真实分段内容"""
    file_id = request.fileId
    
    # 1. 查找文件
    file_path = None
    for ext in ['.pdf', '.docx', '.txt', '.md']:
        path = UPLOAD_DIR / f"{file_id}{ext}"
        if path.exists():
            file_path = path
            break
            
    if not file_path:
        raise HTTPException(status_code=404, detail="Uploaded file not found")

    try:
        sections = []
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.pdf':
            try:
                # Attempt to use Enhanced System (Hierarchical Chunking)
                if str(RAG_DIR) not in sys.path:
                    sys.path.insert(0, str(RAG_DIR))
                from knowledge_base.enhanced_system import PDFProcessor, DotsHierarchicalChunker
                
                print(f"Using Enhanced PDF Processor for {file_path.name}")
                json_doc = PDFProcessor.process(str(file_path))
                
                # Use larger chunks for preview
                chunker = DotsHierarchicalChunker(chunk_size=800, chunk_overlap=50)
                chunks = chunker.chunk(json_doc)

                # Sort by original chunk index to keep natural order
                sorted_chunk_items = [chunks[k] for k in sorted(chunks.keys())]

                for chunk in sorted_chunk_items:
                    # Rehydrate headings text from chunk ids
                    heading_texts = []
                    if chunk.headings:
                        for h_id in chunk.headings:
                            if h_id in chunks:
                                heading_texts.append(chunks[h_id].text)

                    section_title = " > ".join(heading_texts) if heading_texts else f"Page {chunk.page_no}"

                    sections.append({
                        "section": section_title,
                        "content": chunk.text
                    })

                    # Limit preview size
                    if len(sections) >= 50:
                        break
                        
            except Exception as e:
                print(f"Enhanced parsing failed ({e}), falling back to basic loader...")
                try:
                    from langchain_community.document_loaders import PyPDFLoader
                    loader = PyPDFLoader(str(file_path))
                    docs = loader.load()
                except Exception as e2:
                    print(f"PyPDFLoader failed: {e2}")
                    import PyPDF2
                    docs = []
                    with open(file_path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        for i in range(min(len(reader.pages), 5)):
                            page_text = reader.pages[i].extract_text()
                            if page_text:
                                from langchain_core.documents import Document
                                docs.append(Document(page_content=page_text))
                
                # Simple page-based chunking fallback
                for i, doc in enumerate(docs[:5]): 
                    sections.append({
                        "section": f"Page {i+1}",
                        "content": doc.page_content[:2000]
                    })
                
        elif file_ext in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                # 简单按段落分
                paragraphs = content.split('\n\n')
                for i, p in enumerate(paragraphs[:10]):
                    if p.strip():
                        sections.append({
                            "section": f"Paragraph {i+1}",
                            "content": p.strip()
                        })
        
        elif file_ext == '.docx':
            import docx
            doc = docx.Document(file_path)
            for i, p in enumerate(doc.paragraphs[:15]):
                if p.text.strip():
                    sections.append({
                        "section": f"Paragraph {i+1}",
                        "content": p.text.strip()
                    })
        
        if not sections:
            sections = [{"section": "Content", "content": "Failed to extract valid text content."}]

        return {
            "status": "success",
            "message": "Document parsed successfully",
            "data": {
                "title": file_path.name,
                "sections": sections
            }
        }
        
    except Exception as e:
        import traceback
        traceback.print_exc() # 在终端打印详细错误
        raise HTTPException(status_code=500, detail=f"Parsing failed: {str(e)}")
